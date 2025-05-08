import logging
import os
import re
import time
from typing import List
from docx import Document as DocxDocument
from langchain_community.document_loaders.base import BaseLoader
from langchain_core.documents import Document
from processors.keyword_processor import KeywordProcessor


class DocxLoader(BaseLoader):
    """Loader for DOCX files with improved content extraction"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.keyword_processor = KeywordProcessor()

    def load(self) -> List[Document]:
        """Load and process DOCX file with speaker identification"""
        try:
            # Verify file exists and is readable
            if not os.path.exists(self.file_path):
                raise FileNotFoundError(f"File not found: {self.file_path}")

            if os.path.getsize(self.file_path) == 0:
                raise ValueError("File is empty")

            doc = DocxDocument(self.file_path)
            full_text = []
            speaker_paragraphs = []
            current_speaker = None
            current_content = []
            content_found = False

            def add_speaker_block():
                nonlocal content_found
                if current_speaker and current_content:
                    content = '\n'.join(current_content).strip()
                    if content:  # Only add if non-empty content
                        speaker_paragraphs.append({
                            'speaker': current_speaker,
                            'content': content,
                            'keywords': self.keyword_processor.extract_keywords(content)
                        })
                        full_text.append(f"{current_speaker}: {content}")
                        content_found = True

            # Process each paragraph only once
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if not para_text:
                    continue

                # Check for speaker pattern with timestamp
                speaker_match = re.match(r'^(.+?)\s\d+:\d+\s*:', para_text)
                if speaker_match:
                    possible_speaker = speaker_match.group(1).strip()
                    add_speaker_block()
                    current_speaker = possible_speaker
                    content_start = para_text.find(':', speaker_match.end())
                    current_content = [
                        para_text[content_start+1:].strip()] if content_start != -1 else []
                    continue

                # Check for regular speaker pattern
                if ":" in para_text:
                    parts = para_text.split(":", 1)
                    possible_speaker = parts[0].strip()
                    possible_content = parts[1].strip() if len(
                        parts) > 1 else ""

                    if len(possible_speaker.split()) <= 3:  # Simple speaker name heuristic
                        add_speaker_block()
                        current_speaker = possible_speaker
                        current_content = [
                            possible_content] if possible_content else []
                        continue

                # Regular content
                if current_speaker:
                    current_content.append(para_text)
                else:
                    current_speaker = "Document"
                    current_content = [para_text]

            # Add the last speaker block
            add_speaker_block()

            if not content_found:
                logging.warning(
                    f"No readable content found in {self.file_path}")
                return []

            full_content = '\n'.join(full_text).strip()
            if not full_content:
                return []

            # Create metadata
            metadata = {
                "source": self.file_path,
                "filename": os.path.basename(self.file_path),
                "full_document": full_content,
                "is_full_document": True,
                "speaker_blocks": speaker_paragraphs,
                "preprocessed_text": self.keyword_processor.preprocess_text(full_content),
                "keywords": self.keyword_processor.extract_keywords(full_content),
                "file_type": "docx",
                "load_time": time.time()
            }

            return [Document(page_content=full_content, metadata=metadata)]

        except Exception as e:
            logging.error(f"Error loading {self.file_path}: {str(e)}")
            return [Document(
                page_content="",
                metadata={
                    "source": self.file_path,
                    "error": str(e),
                    "status": "failed"
                }
            )]
