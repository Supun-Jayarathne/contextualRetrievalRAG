import hashlib
import json
import os
import re
import time
import logging
from typing import List, Dict, Any, Optional, Tuple
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from loaders.docx_loader import DocxLoader
from processors.keyword_processor import KeywordProcessor
from retrievers.bm25_retriever import BM25Retriever
from retrievers.contextual_retriever import ContextualRetriever
from docx import Document as DocxDocument
from utils import safe_filter_metadata


class ContextualRAGApplication:
    """Enhanced RAG application with contextual embedding"""

    def __init__(self, persist_directory: str = "./chroma_db_docx"):
        """Initialize the Contextual RAG application"""
        self.persist_directory = persist_directory
        self.embeddings = None
        self.llm = None
        self.vectorstore = None
        self.retriever = None
        self.qa_chain = None
        self.contextual_enricher = None
        self.memory = ConversationBufferWindowMemory(k=5, return_messages=True)
        self.bm25_retriever = None
        self.reasoning_steps = 3  # Number of reasoning steps to perform
        self.self_reflection = True  # Whether to enable self-reflection
        self.critique_enabled = True  # Whether to enable answer critique
        self.processed_files = set()  # Track processed file hashes

        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,  # Reduced from 1000
            chunk_overlap=150,  # Reduced from 200
            length_function=len,
            add_start_index=True,
            separators=["\n\n", "\n", " ", ""]  # Explicit separators
        )

        # Load environment variables
        from dotenv import load_dotenv
        load_dotenv()

        try:
            # Initialize components in proper order
            self._initialize_azure_components()
            self._initialize_vectorstore()

            # Now that vectorstore is initialized, create BM25 retriever
            self.bm25_retriever = BM25Retriever(self.vectorstore)
            if self.bm25_retriever.bm25 is None:
                logging.warning(
                    "BM25 retriever disabled - no documents available")

            # Initialize contextual components
            self.contextual_enricher = ContextualRetriever(
                self.llm, self.vectorstore, self.bm25_retriever)
            self._initialize_rag_pipeline()

            logging.info("Contextual RAG application initialized successfully")
        except Exception as e:
            logging.error(f"Failed to initialize BM25 retriever: {str(e)}")
            self.bm25_retriever = None
            raise

    def _initialize_azure_components(self):
        """Initialize Azure OpenAI components"""
        try:
            # Get environment variables with fallbacks
            embedding_deployment = os.getenv(
                "AZURE_OPENAI_EMBEDDING_DEPLOYMENT")
            chat_deployment = os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT")
            api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2023-05-15")
            endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
            api_key = os.getenv("AZURE_OPENAI_API_KEY")

            # Check for required configuration
            if not all([embedding_deployment, chat_deployment, endpoint, api_key]):
                missing = []
                if not embedding_deployment:
                    missing.append("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")
                if not chat_deployment:
                    missing.append("AZURE_OPENAI_CHAT_DEPLOYMENT")
                if not endpoint:
                    missing.append("AZURE_OPENAI_ENDPOINT")
                if not api_key:
                    missing.append("AZURE_OPENAI_API_KEY")

                raise ValueError(
                    f"Missing required Azure OpenAI configuration: {', '.join(missing)}")

            logging.info(
                f"Initializing Azure OpenAI with: Embedding deployment: {embedding_deployment}, Chat deployment: {chat_deployment}")

            # Initialize embeddings
            try:
                self.embeddings = AzureOpenAIEmbeddings(
                    azure_deployment=embedding_deployment,
                    openai_api_version=api_version,
                    azure_endpoint=endpoint,
                    api_key=api_key
                )
                # Test the embeddings with a simple query to catch errors early
                test_result = self.embeddings.embed_query("Test embedding")
                if not test_result or len(test_result) < 1:
                    raise ValueError(
                        "Embedding test failed - returned empty result")
                logging.info(
                    f"Embeddings successfully initialized and tested ({len(test_result)} dimensions)")
            except Exception as e:
                error_msg = str(e)
                if "OperationNotSupported" in error_msg:
                    raise ValueError(
                        f"The model in deployment '{embedding_deployment}' doesn't support embeddings. "
                        f"Please use a text-embedding model like 'text-embedding-ada-002'. Error: {error_msg}"
                    )
                else:
                    raise ValueError(
                        f"Failed to initialize embeddings: {error_msg}")

            # Initialize LLM
            try:
                self.llm = AzureChatOpenAI(
                    azure_deployment=chat_deployment,
                    openai_api_version=api_version,
                    azure_endpoint=endpoint,
                    api_key=api_key,
                    temperature=0.1,  # Low temperature for factual responses
                    max_tokens=1000
                )
                logging.info(f"Chat model successfully initialized")
            except Exception as e:
                raise ValueError(f"Failed to initialize chat model: {str(e)}")

            logging.info("Azure OpenAI components initialized successfully")
        except Exception as e:
            logging.error(f"Failed to initialize Azure components: {str(e)}")
            raise

    def _initialize_contextual_enricher(self):
        """Initialize the contextual enrichment component"""
        try:
            self.contextual_enricher = ContextualRetriever(
                self.llm, self.vectorstore)
            logging.info("Contextual enricher initialized successfully")
        except Exception as e:
            logging.error(
                f"Failed to initialize contextual enricher: {str(e)}")
            raise

    def _initialize_vectorstore(self):
        """Initialize or load the vector store"""
        try:
            # Check if persist directory exists
            if os.path.exists(self.persist_directory) and os.listdir(self.persist_directory):
                # Load existing vector store
                logging.info(
                    f"Loading vector store from {self.persist_directory}")
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                logging.info(
                    f"Loaded existing vector store with {self.vectorstore._collection.count()} documents")
            else:
                # Create directory if it doesn't exist
                os.makedirs(self.persist_directory, exist_ok=True)

                # Create new vector store
                logging.info("Creating new vector store")
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                self.vectorstore.persist()
                logging.info("Created new empty vector store")

            # Verify the vectorstore is properly initialized
            if not hasattr(self.vectorstore, '_collection'):
                raise RuntimeError(
                    "Vector store failed to initialize properly - missing _collection attribute")

        except Exception as e:
            logging.error(f"Failed to initialize vector store: {str(e)}")
            self.vectorstore = None  # Ensure it's None if initialization fails
            raise

    def _initialize_rag_pipeline(self):
        """Initialize RAG pipeline with hybrid retrieval and auto-reasoning"""
        try:
            # Enhanced prompt with reasoning structure
            prompt_template = """You are a document intelligence assistant with advanced reasoning capabilities. 
            Follow this structured approach to answer questions:

            --- REASONING PROCESS ---
            1. CONTEXT ANALYSIS:
            {contextual_analysis}

            2. RELEVANT CONTENT:
            {context}

            3. CHAT HISTORY:
            {chat_history}

            4. QUESTION ANALYSIS:
            - Type: {question_type}
            - Key Entities: {key_entities}
            - Potential Ambiguities: {ambiguities}

            5. STEP-BY-STEP REASONING:
            {reasoning_steps}

            6. FINAL ANSWER:
            {final_answer}

            Guidelines:
            1. First analyze the question type and requirements
            2. Examine all contextual relevance information
            3. Perform {reasoning_steps} explicit reasoning steps
            4. Cross-validate information across sources
            5. Consider alternative interpretations
            6. If uncertain, specify what information is missing
            7. For complex questions, break down into sub-questions
            8. Mark assumptions and confidence levels
            """

            PROMPT = PromptTemplate(
                template=prompt_template,
                input_variables=[
                    "contextual_analysis",
                    "context",
                    "question",
                    "chat_history",
                    "question_type",
                    "key_entities",
                    "ambiguities",
                    "reasoning_steps",
                    "final_answer"
                ]
            )

            # Create analysis components
            def analyze_question(question: str) -> dict:
                """Robust question analysis with multiple JSON parsing fallbacks"""
                analysis_prompt = """ONLY output a JSON object with these exact keys:
                                {
                                    "question_type": ("factual", "interpretive", "comparative", "procedural", "opinion", "unknown"),
                                    "key_entities": ["list", "of", "entities"],
                                    "ambiguities": ["list", "of", "ambiguities"]
                                }

                                Question: {question}

                                Rules:
                                1. Output MUST be parseable by json.loads()
                                2. NO additional text before/after the JSON
                                3. ONLY use double quotes
                                4. NO trailing commas
                                5. Keys MUST match exactly
                                """

                default_response = {
                    "question_type": "unknown",
                    "key_entities": [],
                    "ambiguities": []
                }

                try:
                    # Get response with strict parameters
                    response = self.llm.invoke(
                        PromptTemplate.from_template(
                            analysis_prompt).format(question=question),
                        temperature=0,
                        max_tokens=200
                    ).content

                    # Clean response through multiple fallback strategies
                    json_str = self._extract_json(response)
                    if not json_str:
                        return default_response

                    # Parse with validation
                    result = json.loads(json_str)
                    if not all(k in result for k in ["question_type", "key_entities", "ambiguities"]):
                        return default_response

                    # Normalize types
                    result["key_entities"] = list(
                        result.get("key_entities", []))
                    result["ambiguities"] = list(result.get("ambiguities", []))
                    return result

                except Exception as e:
                    logging.warning(f"Question analysis failed: {str(e)}")
                    return default_response

            def generate_reasoning_steps(question: str, context: str, question_analysis: dict) -> str:
                """Generate explicit reasoning steps"""
                reasoning_prompt = """Generate {steps} reasoning steps to answer this question:
                Question: {question}
                Type: {question_type}
                Context: {context}
                
                For each step, include:
                - Purpose of the step
                - Method of analysis
                - Expected outcome
                """
                try:
                    reasoning = self.llm.invoke(
                        PromptTemplate.from_template(reasoning_prompt).format(
                            steps=self.reasoning_steps,
                            question=question,
                            question_type=question_analysis.get(
                                "question_type", "unknown"),
                            context=context
                        )
                    )
                    return reasoning.content
                except Exception as e:
                    logging.warning(
                        f"Reasoning steps generation failed: {str(e)}")
                    return "Could not generate explicit reasoning steps."

            def critique_answer(answer: str, context: str) -> str:
                """Critique and improve the answer"""
                if not self.critique_enabled:
                    return answer

                critique_prompt = """Critique and improve this answer based on the context:
                Context: {context}
                Initial Answer: {answer}
                
                Provide:
                1. Potential weaknesses
                2. Missing information
                3. Improved version
                """
                try:
                    critique = self.llm.invoke(
                        PromptTemplate.from_template(critique_prompt).format(
                            context=context,
                            answer=answer
                        )
                    )
                    return critique.content
                except Exception as e:
                    logging.warning(f"Answer critique failed: {str(e)}")
                    return answer

            # Enhanced hybrid retriever with reasoning
            def hybrid_retriever_with_reasoning(input_dict: Dict[str, Any]) -> Dict[str, Any]:
                """Enhanced retrieval with reasoning components"""
                try:
                    question = input_dict["question"]

                    # Get documents and contextual analysis
                    docs = self.contextual_enricher.hybrid_search(
                        question, k=10)

                    # Extract only relevant sentences from each document
                    def extract_relevant_sentences(content: str, question: str) -> str:
                        # Simple implementation - in practice you might use more sophisticated NLP
                        sentences = [s.strip()
                                     for s in content.split('.') if s.strip()]
                        relevant = []
                        question_lower = question.lower()
                        for sentence in sentences:
                            if any(word.lower() in sentence.lower() for word in question.split()):
                                relevant.append(f'"{sentence.strip()}."')
                                if len(relevant) >= 3:  # Limit to 3 most relevant sentences
                                    break
                        return ' '.join(relevant) if relevant else '"Most relevant content not found in this document."'

                    contextual_analysis = "\n\n".join(
                        f"Document {i+1} (Relevance: {doc.metadata.get('contextual_score', 0):.0%}, Method: {doc.metadata.get('retrieval_method', 'unknown')}):\n"
                        f"Explanation: {doc.metadata.get('contextual_explanation', '')}\n"
                        f"Connections: {', '.join(doc.metadata.get('contextual_connections', []))}"
                        for i, doc in enumerate(docs)
                    )

                    # Show only relevant sentences with quotes instead of full content
                    context_content = "\n\n".join(
                        f"Document {i+1} relevant content:\n{extract_relevant_sentences(doc.page_content, question)}"
                        for i, doc in enumerate(docs))

                    # Question analysis
                    question_analysis = analyze_question(question)

                    # Generate reasoning steps
                    reasoning_steps = generate_reasoning_steps(
                        question,
                        context_content,
                        question_analysis
                    )

                    # Generate initial answer
                    initial_answer = self.llm.invoke(
                        PromptTemplate.from_template(
                            "Answer this question based on the relevant context:\nQuestion: {question}\nRelevant Context: {context}"
                        ).format(
                            question=question,
                            context=context_content
                        )
                    ).content

                    # Critique and improve answer
                    final_answer = critique_answer(
                        initial_answer, context_content)

                    return {
                        "contextual_analysis": contextual_analysis,
                        "context": context_content,
                        "question": question,
                        "chat_history": input_dict.get("chat_history", ""),
                        "question_type": question_analysis.get("question_type", "unknown"),
                        "key_entities": ", ".join(question_analysis.get("key_entities", [])),
                        "ambiguities": ", ".join(question_analysis.get("ambiguities", [])),
                        "reasoning_steps": reasoning_steps,
                        "final_answer": final_answer
                    }
                except Exception as e:
                    logging.error(f"Error in reasoning retriever: {str(e)}")
                    return {
                        "contextual_analysis": "Error in reasoning process",
                        "context": "",
                        "question": question,
                        "chat_history": input_dict.get("chat_history", ""),
                        "question_type": "unknown",
                        "key_entities": "",
                        "ambiguities": "",
                        "reasoning_steps": "",
                        "final_answer": f"Error generating answer: {str(e)}"
                    }

            # Build the full chain with reasoning
            self.qa_chain = (
                RunnablePassthrough()
                | hybrid_retriever_with_reasoning
                | PROMPT
                | self.llm
                | StrOutputParser()
            )

            logging.info("Enhanced RAG pipeline with reasoning initialized")

        except Exception as e:
            logging.error(f"Error initializing reasoning pipeline: {str(e)}")
            raise

    def _extract_json(self, text: str) -> Optional[str]:
        """Multiple strategies to extract valid JSON from LLM output"""
        text = text.strip()

        # Strategy 1: Direct parse if clean
        try:
            json.loads(text)
            return text
        except json.JSONDecodeError:
            pass

        # Strategy 2: Remove common non-JSON prefixes
        patterns = [
            r'^[^\{]*',  # Anything before first {
            r'^.*?\{(.*)\}.*$',  # Extract between curly braces
            r'```json(.*)```'  # Extract from markdown code block
        ]

        for pattern in patterns:
            try:
                match = re.search(pattern, text, re.DOTALL)
                if match:
                    candidate = match.group(1) if pattern.startswith(
                        '^.*?') else match.group(0)
                    candidate = candidate.strip()
                    if candidate.startswith('{'):
                        json.loads(candidate)
                        return candidate
            except:
                continue

        # Strategy 3: Repair attempts
        repair_attempts = [
            lambda x: x.replace("'", '"'),  # Single to double quotes
            # Remove trailing commas
            lambda x: re.sub(r',\s*([}\]])', r'\1', x),
            # Remove empty elements
            lambda x: re.sub(r'([{\[,])\s*([}\]])', r'\1\2', x),
            lambda x: x.replace('\n', '').replace(
                '\t', '')  # Remove whitespace
        ]

        for attempt in repair_attempts:
            try:
                repaired = attempt(text)
                json.loads(repaired)
                return repaired
            except:
                continue

        return None

    def _get_chat_history(self) -> str:
        """Get formatted chat history from memory"""
        messages = self.memory.load_memory_variables({}).get("history", [])
        formatted_messages = []

        for message in messages:
            if hasattr(message, "content"):
                role = "Human" if message.type == "human" else "Assistant"
                formatted_messages.append(f"{role}: {message.content}")

        return "\n".join(formatted_messages)

    def _get_file_hash(self, file_content: bytes) -> str:
        """Generate a consistent hash for file content"""
        return hashlib.md5(file_content).hexdigest()

    def add_uploaded_files(self, files) -> str:
        """Process and add uploaded files to the vector store with hash-based deduplication"""
        try:
            if not files:
                return "No files provided"

            documents = []
            stats = {
                'processed_files': 0,
                'skipped_files': 0,
                'new_chunks': 0,
                'failed_files': 0
            }

            for file in files:
                file_name = "unknown"
                try:
                    # Handle different file input types
                    if isinstance(file, tuple) and len(file) == 2:
                        file_name, file_content = file
                        file_bytes = file_content
                    elif hasattr(file, 'name') and hasattr(file, 'read'):
                        file_name = file.name
                        file.seek(0)
                        file_bytes = file.read()
                    else:
                        file_name = f"upload_{time.time()}.docx"
                        file_bytes = file

                    # Validate file type
                    if not file_name.lower().endswith('.docx'):
                        logging.warning(f"Skipping non-DOCX file: {file_name}")
                        stats['skipped_files'] += 1
                        continue

                    # Generate content-based hash
                    file_hash = hashlib.md5(file_bytes).hexdigest()

                    # Check if file already exists
                    if self._is_file_processed(file_hash):
                        logging.info(
                            f"Skipping already processed file: {file_name}")
                        stats['skipped_files'] += 1
                        continue

                    # Create temp directory if it doesn't exist
                    temp_dir = "./temp_uploads"
                    os.makedirs(temp_dir, exist_ok=True)

                    # Create temp file path
                    clean_name = re.sub(
                        r'[^\w.-]', '_', os.path.basename(file_name))
                    temp_path = os.path.join(
                        temp_dir, f"upload_{file_hash[:8]}_{clean_name}")

                    # Write to temp file
                    with open(temp_path, "wb") as f:
                        f.write(file_bytes)

                    # Verify file is valid DOCX
                    try:
                        doc = DocxDocument(temp_path)
                        if not doc.paragraphs:
                            raise ValueError("No paragraphs found in document")
                    except Exception as e:
                        logging.warning(
                            f"Invalid DOCX file {file_name}: {str(e)}")
                        os.remove(temp_path)
                        stats['failed_files'] += 1
                        continue

                    # Load document
                    loader = DocxLoader(temp_path)
                    loaded_docs = loader.load()
                    valid_docs = []

                    for doc in loaded_docs:
                        if doc.page_content.strip():
                            filtered_doc = safe_filter_metadata(doc)
                            if isinstance(filtered_doc, Document) and filtered_doc.page_content.strip():
                                # Add file hash to metadata
                                filtered_doc.metadata['file_hash'] = file_hash
                                valid_docs.append(filtered_doc)

                    if valid_docs:
                        documents.extend(valid_docs)
                        stats['processed_files'] += 1
                        self._mark_file_processed(
                            file_hash)  # Mark as processed

                    # Clean up temp file
                    os.remove(temp_path)

                except Exception as e:
                    logging.error(
                        f"Error processing file {file_name}: {str(e)}")
                    stats['failed_files'] += 1
                    continue

            if not documents:
                return (
                    f"No valid documents extracted. "
                    f"Skipped {stats['skipped_files']} duplicates, "
                    f"{stats['failed_files']} failed"
                )

            # Process documents through the standardized pipeline
            processed_chunks = self._process_documents(documents)

            if not processed_chunks:
                return "No valid chunks created after processing"

            # Filter out chunks that already exist in the database
            existing_hashes = self._get_existing_chunk_hashes()
            new_chunks = [
                chunk for chunk in processed_chunks
                if chunk.metadata.get('file_hash') not in existing_hashes
            ]

            if not new_chunks:
                return (
                    f"All chunks already exist in database. "
                    f"Processed {stats['processed_files']} files, "
                    f"skipped {stats['skipped_files']} duplicates"
                )

            # Store only new chunks
            try:
                self.vectorstore.add_documents(new_chunks)
                self.vectorstore.persist()

                # Reinitialize retrievers if needed
                if self.bm25_retriever:
                    self.bm25_retriever._initialize_bm25()

                stats['new_chunks'] = len(new_chunks)
                return (
                    f"Added {stats['new_chunks']} new chunks from {stats['processed_files']} files. "
                    f"Skipped {stats['skipped_files']} duplicates. "
                    f"{stats['failed_files']} files failed processing."
                )

            except Exception as e:
                logging.error(f"Error storing documents: {str(e)}")
                return f"Error storing documents: {str(e)}"

        except Exception as e:
            logging.error(f"Error in add_uploaded_files: {str(e)}")
            return f"Error processing files: {str(e)}"

    def _is_file_processed(self, file_hash: str) -> bool:
        """Check if file has already been processed"""
        # Implement your persistence method here (could be in-memory set, DB, etc.)
        return file_hash in self.processed_files

    def _mark_file_processed(self, file_hash: str):
        """Mark a file as processed"""
        self.processed_files.add(file_hash)
        # Add persistence logic if needed

    def _get_existing_chunk_hashes(self) -> set:
        """Get hashes of all existing chunks in the vector store"""
        try:
            collection = self.vectorstore._collection
            if collection.count() == 0:
                return set()

            return {
                m.get('file_hash')
                for m in collection.get()['metadatas']
                if m.get('file_hash')
            }
        except Exception as e:
            logging.error(f"Error getting existing hashes: {str(e)}")
            return set()

    def _process_documents(self, documents: List[Document]) -> List[Document]:
        """Process documents with proper metadata filtering and contextual enrichment"""
        processed_chunks = []
        keyword_processor = KeywordProcessor()
        processed_hashes = set()

        def get_chunk_hash(chunk):
            content_hash = hashlib.md5(chunk.page_content.encode()).hexdigest()
            metadata_hash = hashlib.md5(json.dumps(
                chunk.metadata, sort_keys=True).encode()).hexdigest()
            return f"{content_hash}:{metadata_hash}"

        for doc in documents:
            if not isinstance(doc, Document) or not doc.page_content.strip():
                continue

            try:
                filtered_doc = safe_filter_metadata(doc)
                metadata = filtered_doc.metadata.copy()

                metadata.update({
                    "source": metadata.get("source", "unknown"),
                    "filename": metadata.get("filename", "unknown"),
                    "full_document": metadata.get("full_document", filtered_doc.page_content),
                    "processing_time": time.time()
                })

                chunks = self.text_splitter.split_documents([filtered_doc])

                for chunk in chunks:
                    if not isinstance(chunk, Document) or not chunk.page_content.strip() or len(chunk.page_content.strip()) < 50:
                        continue

                    try:
                        chunk_metadata = metadata.copy()
                        content = chunk.page_content

                        keywords = keyword_processor.extract_keywords(content)
                        chunk_metadata["keywords"] = ", ".join(
                            sorted(keywords)) if keywords else ""
                        preprocessed = keyword_processor.preprocess_text(
                            content)
                        chunk_metadata["preprocessed_text"] = preprocessed if isinstance(
                            preprocessed, str) else ""

                        chunk_metadata.update({
                            "original_content": content,
                            "is_full_document": False,
                            "chunk_length": len(content),
                            "word_count": len(content.split())
                        })

                        filtered_metadata = safe_filter_metadata(
                            chunk_metadata)
                        final_metadata = filtered_metadata if isinstance(
                            filtered_metadata, dict) else filtered_metadata.metadata

                        base_chunk = Document(
                            page_content=content,
                            metadata=final_metadata
                        )

                        # Try to enrich, otherwise use base chunk
                        try:
                            final_chunk = self.contextual_enricher.enrich_chunk(
                                base_chunk) if hasattr(self, 'contextual_enricher') else base_chunk
                            if not isinstance(final_chunk, Document) or not final_chunk.page_content.strip():
                                final_chunk = base_chunk
                        except Exception:
                            final_chunk = base_chunk

                        # Deduplication check
                        chunk_hash = get_chunk_hash(final_chunk)
                        if chunk_hash not in processed_hashes:
                            processed_chunks.append(final_chunk)
                            processed_hashes.add(chunk_hash)

                    except Exception as chunk_error:
                        logging.error(
                            f"Error processing chunk: {str(chunk_error)}")
                        continue

            except Exception as doc_error:
                logging.error(f"Error processing document: {str(doc_error)}")
                continue

        return processed_chunks

    def query(self, question: str, add_to_memory: bool = True) -> Dict[str, Any]:
        """Query the contextual RAG pipeline"""
        try:
            # Check if components are properly initialized
            if self.vectorstore is None:
                return {
                    "result": "Document system is not properly initialized. Please try resetting the document store.",
                    "source_documents": [],
                    "query_time": None
                }

            if not hasattr(self.vectorstore, '_collection'):
                return {
                    "result": "Vector store is not properly configured. Please try resetting the document store.",
                    "source_documents": [],
                    "query_time": None
                }

            # Check if the vector store is empty
            try:
                if self.vectorstore._collection.count() == 0:
                    return {
                        "result": "No documents have been added to the system. Please upload some documents first.",
                        "source_documents": [],
                        "query_time": None
                    }
            except Exception as e:
                logging.error(f"Error checking vector store count: {str(e)}")
                return {
                    "result": "Error checking document store status.",
                    "source_documents": [],
                    "query_time": None
                }

            # Get start time
            start_time = time.time()

            # Get the answer from the QA chain
            try:
                answer = self.qa_chain.invoke({
                    "question": question,
                    "chat_history": self._get_chat_history()
                })
            except Exception as e:
                logging.error(f"Error in QA chain: {str(e)}")
                answer = f"Error generating answer: {str(e)}"

            # Get the source documents with contextual information
            try:
                source_docs = []
                if self.contextual_enricher is not None:
                    source_docs = self.contextual_enricher.contextual_search(
                        question)
            except Exception as e:
                logging.error(f"Error retrieving source documents: {str(e)}")
                source_docs = []

            # Calculate query time
            query_time = time.time() - start_time

            # Add to memory if requested
            if add_to_memory and self.memory is not None:
                try:
                    self.memory.save_context(
                        {"input": question},
                        {"output": answer}
                    )
                except Exception as e:
                    logging.error(f"Error saving to memory: {str(e)}")

            return {
                "result": answer,
                "source_documents": source_docs,
                "query_time": f"{query_time:.2f} seconds"
            }
        except Exception as e:
            logging.error(
                f"Unexpected error in query: {str(e)}", exc_info=True)
            return {
                "result": f"Unexpected error processing your request: {str(e)}",
                "source_documents": [],
                "query_time": None
            }

    def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about the documents in the vector store"""
        try:
            # Get document count
            doc_count = self.vectorstore._collection.count()

            # Get unique sources
            if doc_count > 0:
                docs = self.vectorstore._collection.get()
                metadatas = docs["metadatas"]
                unique_sources = set()

                # Collect topic statistics
                all_topics = []
                all_entities = []

                for metadata in metadatas:
                    if "source" in metadata:
                        unique_sources.add(metadata["source"])

                    # Collect topics and entities for statistics
                    if "topics" in metadata and isinstance(metadata["topics"], list):
                        all_topics.extend(metadata["topics"])
                    if "entities" in metadata and isinstance(metadata["entities"], list):
                        all_entities.extend(metadata["entities"])

                # Get top topics and entities
                top_topics = self._get_top_items(all_topics, 10)
                top_entities = self._get_top_items(all_entities, 10)

            else:
                unique_sources = set()
                top_topics = []
                top_entities = []

            return {
                "document_chunks": doc_count,
                "unique_files": len(unique_sources),
                "files": list(unique_sources),
                "top_topics": top_topics,
                "top_entities": top_entities
            }
        except Exception as e:
            logging.error(f"Error getting document stats: {str(e)}")
            return {
                "document_chunks": 0,
                "unique_files": 0,
                "files": [],
                "top_topics": [],
                "top_entities": []
            }

    def _get_top_items(self, items: List[str], limit: int = 10) -> List[Tuple[str, int]]:
        """Get top items by frequency"""
        from collections import Counter
        if not items:
            return []

        counter = Counter(items)
        return counter.most_common(limit)

    def clear_memory(self) -> None:
        """Clear the conversation memory"""
        self.memory.clear()
        logging.info("Conversation memory cleared")

    def reset_vectorstore(self) -> str:
        """Reset the vector store, removing all documents"""
        try:
            # Close existing connections first
            if self.vectorstore:
                try:
                    self.vectorstore._client = None
                    self.vectorstore._collection = None
                except Exception as e:
                    logging.warning(f"Error cleaning up vectorstore: {str(e)}")

            # Delete the persist directory
            import shutil
            if os.path.exists(self.persist_directory):
                for _ in range(3):  # Retry up to 3 times
                    try:
                        shutil.rmtree(self.persist_directory)
                        break
                    except PermissionError as e:
                        logging.warning(f"Retrying delete: {str(e)}")
                        time.sleep(0.5)  # Wait briefly before retrying
                    except Exception as e:
                        logging.error(f"Error deleting directory: {str(e)}")
                        raise

            # Reinitialize the vector store
            self._initialize_vectorstore()

            logging.info("Vector store reset successfully")
            return "✅ Vector store reset successfully. All documents have been removed."
        except Exception as e:
            logging.error(f"Error resetting vector store: {str(e)}")
            return f"❌ Error resetting vector store: {str(e)}"
