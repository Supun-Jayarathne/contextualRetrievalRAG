import os
import gradio as gr
import pandas as pd
from dotenv import load_dotenv
from typing import List, Iterator, Optional, Dict, Any, Tuple, Union
from langchain_community.document_loaders.base import BaseLoader
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from docx import Document as DocxDocument
from rank_bm25 import BM25Okapi
from typing import List, Dict, Any, Optional
import numpy as np
from io import BytesIO
import logging
import time
import hashlib
import json
from tqdm.auto import tqdm
from collections import Counter
import re
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
import nltk

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("rag_app.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Custom function to safely filter complex metadata


def safe_filter_metadata(doc: Union[Document, dict]) -> Union[Document, dict]:
    """Safely filter metadata to only include simple types that ChromaDB can handle.
    Handles both Document objects and raw dictionaries."""
    try:
        if isinstance(doc, Document):
            # Handle Document objects
            metadata = getattr(doc, 'metadata', {})
            if not isinstance(metadata, dict):
                metadata = {}

            filtered_metadata = {}
            for key, value in metadata.items():
                # Skip None values
                if value is None:
                    continue

                # Handle allowed simple types
                if isinstance(value, (str, int, float, bool)):
                    filtered_metadata[key] = value
                # Convert lists to comma-separated strings
                elif isinstance(value, list):
                    filtered_metadata[key] = ", ".join(
                        str(item) for item in value
                        if isinstance(item, (str, int, float, bool))
                    )
                # Convert other types to strings
                else:
                    try:
                        filtered_metadata[key] = str(value)
                    except:
                        continue

            return Document(
                page_content=doc.page_content,
                metadata=filtered_metadata
            )

        elif isinstance(doc, dict):
            # Handle dictionary inputs
            filtered = {}
            for key, value in doc.items():
                if value is None:
                    continue
                if isinstance(value, (str, int, float, bool)):
                    filtered[key] = value
                elif isinstance(value, list):
                    filtered[key] = ", ".join(
                        str(item) for item in value
                        if isinstance(item, (str, int, float, bool))
                    )
                else:
                    try:
                        filtered[key] = str(value)
                    except:
                        continue
            return filtered

        raise ValueError("Input must be a Document object or dictionary")

    except Exception as e:
        logger.error(f"Error filtering metadata: {str(e)}")
        # Fallback to minimal metadata
        if isinstance(doc, Document):
            return Document(
                page_content=doc.page_content,
                metadata={"source": "unknown",
                          "error": "metadata_filter_failed"}
            )
        return {"source": "unknown", "error": "metadata_filter_failed"}


class DocxLoader(BaseLoader):
    """Loader for DOCX files with improved content extraction"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.keyword_processor = KeywordProcessor()

    def load(self) -> List[Document]:
        """Load and process DOCX file with speaker identification"""
        try:
            # Verify file exists and is readable
            if not os.path.exists(self.file_path):
                raise FileNotFoundError(f"File not found: {self.file_path}")

            if os.path.getsize(self.file_path) == 0:
                raise ValueError("File is empty")

            doc = DocxDocument(self.file_path)
            full_text = []
            speaker_paragraphs = []
            current_speaker = None
            current_content = []
            content_found = False

            def add_speaker_block():
                nonlocal content_found
                if current_speaker and current_content:
                    content = '\n'.join(current_content).strip()
                    if content:  # Only add if non-empty content
                        speaker_paragraphs.append({
                            'speaker': current_speaker,
                            'content': content,
                            'keywords': self.keyword_processor.extract_keywords(content)
                        })
                        full_text.append(f"{current_speaker}: {content}")
                        content_found = True

            # Process each paragraph only once
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if not para_text:
                    continue

                # Check for speaker pattern with timestamp
                speaker_match = re.match(r'^(.+?)\s\d+:\d+\s*:', para_text)
                if speaker_match:
                    possible_speaker = speaker_match.group(1).strip()
                    add_speaker_block()
                    current_speaker = possible_speaker
                    content_start = para_text.find(':', speaker_match.end())
                    current_content = [
                        para_text[content_start+1:].strip()] if content_start != -1 else []
                    continue

                # Check for regular speaker pattern
                if ":" in para_text:
                    parts = para_text.split(":", 1)
                    possible_speaker = parts[0].strip()
                    possible_content = parts[1].strip() if len(
                        parts) > 1 else ""

                    if len(possible_speaker.split()) <= 3:  # Simple speaker name heuristic
                        add_speaker_block()
                        current_speaker = possible_speaker
                        current_content = [
                            possible_content] if possible_content else []
                        continue

                # Regular content
                if current_speaker:
                    current_content.append(para_text)
                else:
                    current_speaker = "Document"
                    current_content = [para_text]

            # Add the last speaker block
            add_speaker_block()

            if not content_found:
                logger.warning(
                    f"No readable content found in {self.file_path}")
                return []

            full_content = '\n'.join(full_text).strip()
            if not full_content:
                return []

            # Create metadata
            metadata = {
                "source": self.file_path,
                "filename": os.path.basename(self.file_path),
                "full_document": full_content,
                "is_full_document": True,
                "speaker_blocks": speaker_paragraphs,
                "preprocessed_text": self.keyword_processor.preprocess_text(full_content),
                "keywords": self.keyword_processor.extract_keywords(full_content),
                "file_type": "docx",
                "load_time": time.time()
            }

            return [Document(page_content=full_content, metadata=metadata)]

        except Exception as e:
            logger.error(f"Error loading {self.file_path}: {str(e)}")
            return [Document(
                page_content="",
                metadata={
                    "source": self.file_path,
                    "error": str(e),
                    "status": "failed"
                }
            )]


class BM25Retriever:
    """Enhanced BM25 retriever with proper empty collection handling"""

    def __init__(self, vectorstore: Chroma):
        self.vectorstore = vectorstore
        self.bm25 = None
        self.documents = []
        self.keyword_processor = KeywordProcessor()
        self._initialize_bm25()

    def _initialize_bm25(self):
        """Initialize BM25 with proper error handling"""
        try:
            collection = self.vectorstore._collection.get()
            self.documents = []

            # Check if we have any documents
            if not collection.get('documents'):
                logger.warning(
                    "Vector store contains no documents - BM25 will be disabled")
                self.bm25 = None
                return

            for content, meta in zip(collection['documents'], collection['metadatas']):
                # Skip empty documents
                if not content.strip():
                    continue

                # Combine keywords and preprocessed text
                keywords = meta.get('keywords', [])
                preprocessed = meta.get('preprocessed_text',
                                        self.keyword_processor.preprocess_text(content))

                # Create hybrid document representation
                doc_representation = f"{' '.join(keywords)} {preprocessed}"
                self.documents.append(doc_representation)

            # Check if we ended up with any valid documents
            if not self.documents:
                logger.warning(
                    "No valid documents found for BM25 initialization")
                self.bm25 = None
                return

            tokenized_docs = [doc.split() for doc in self.documents]
            self.bm25 = BM25Okapi(tokenized_docs)
            logger.info(
                f"BM25 retriever initialized with {len(self.documents)} documents")

        except Exception as e:
            logger.error(f"Error initializing BM25: {str(e)}")
            self.bm25 = None  # Ensure it's None if initialization fails
            raise

    def get_hybrid_query(self, query: str) -> str:
        """Enhance the query with extracted keywords"""
        if not query.strip():
            return ""

        query_keywords = self.keyword_processor.extract_keywords(query)
        query_preprocessed = self.keyword_processor.preprocess_text(query)
        return f"{' '.join(query_keywords)} {query_preprocessed}"

    def retrieve(self, query: str, k: int = 5) -> List[Document]:
        """Retrieve documents with proper empty collection handling"""
        try:
            if not self.bm25 or not self.documents:
                return []

            if not query.strip():
                return []

            # Enhance the query
            enhanced_query = self.get_hybrid_query(query)

            # Get top documents with BM25
            tokenized_query = enhanced_query.split()
            doc_scores = self.bm25.get_scores(tokenized_query)
            top_indices = np.argsort(doc_scores)[::-1][:k]

            # Get corresponding documents from vector store
            collection = self.vectorstore._collection.get()
            retrieved_docs = []

            for idx in top_indices:
                if idx < len(collection['ids']):
                    doc_id = collection['ids'][idx]
                    doc_content = collection['documents'][idx]
                    doc_metadata = collection['metadatas'][idx] if 'metadatas' in collection else {
                    }

                    # Add BM25 score to metadata
                    metadata = doc_metadata.copy()
                    metadata['bm25_score'] = float(doc_scores[idx])
                    metadata['bm25_query'] = enhanced_query

                    retrieved_docs.append(Document(
                        page_content=doc_content,
                        metadata=metadata
                    ))

            return retrieved_docs

        except Exception as e:
            logger.error(f"Error in BM25 retrieval: {str(e)}")
            return []

    # Download NLTK resources if not already present
    try:
        nltk.data.find('tokenizers/punkt')
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('punkt')
        nltk.download('stopwords')


class KeywordProcessor:
    """Extracts meaningful keywords representing main ideas, concepts, and themes from text"""

    def __init__(self):
        # Basic stopwords
        self.stop_words = set(stopwords.words('english'))
        self.stemmer = PorterStemmer()
        self.min_word_length = 3
        self.max_keywords = 20

        # POS tags that typically represent meaningful concepts
        self.concept_pos_tags = {
            'NN', 'NNS',  # Nouns
            'NNP', 'NNPS',  # Proper nouns
            'JJ',  # Adjectives (often part of key concepts)
        }

    def extract_keywords(self, text: str, use_stemming: bool = False) -> List[str]:
        """
        Extract meaningful keywords (nouns, abbreviations, domain terms) from text

        Args:
            text: The input text to extract keywords from
            use_stemming: Whether to apply stemming to keywords (default: False)

        Returns:
            List of extracted keywords
        """
        if not text or len(text.strip()) == 0:
            return []

        try:
            # Extract abbreviations (capitalized terms and acronyms)
            abbreviations = set(re.findall(r'\b[A-Z][A-Z0-9]+\b', text))

            # Extract potential domain-specific terms (including hyphenated terms)
            domain_terms = set(re.findall(r'\b[A-Za-z]+-?[A-Za-z]+\b', text))

            # Tokenize text for normal processing
            words = word_tokenize(text)

            # Apply POS tagging to identify nouns and important word types
            try:
                from nltk import pos_tag
                tagged_words = pos_tag(words)

                # Get words with concept tags (nouns, proper nouns, etc.)
                concept_words = [word.lower() for word, tag in tagged_words
                                 if tag in self.concept_pos_tags and
                                 len(word) >= self.min_word_length and
                                 word.lower() not in self.stop_words]
            except:
                # Fallback if NLTK POS tagging fails
                concept_words = [w.lower() for w in words if len(w) >= self.min_word_length and
                                 w.lower() not in self.stop_words]

            # Combine all potential keywords, prioritizing abbreviations and domain terms
            all_keywords = list(abbreviations) + concept_words + [t.lower() for t in domain_terms
                                                                  if len(t) >= self.min_word_length]

            # Count frequencies and get most common
            keyword_counts = Counter(all_keywords)
            keywords = [w for w, _ in keyword_counts.most_common(
                self.max_keywords)]

            # Apply stemming if requested
            if use_stemming and keywords:
                return [self.stemmer.stem(kw) for kw in keywords]

            return keywords

        except Exception as e:
            print(f"Error extracting keywords: {str(e)}")
            # Last resort - extract nouns with regex
            potential_nouns = re.findall(r'\b[A-Z]?[a-z]+\b', text)
            return [w for w in potential_nouns if len(w) >= self.min_word_length
                    and w.lower() not in self.stop_words][:self.max_keywords]

    def extract_key_phrases(self, text: str, max_phrases: int = 10) -> List[str]:
        """
        Extract multi-word phrases that represent concepts

        Args:
            text: Input text
            max_phrases: Maximum number of phrases to return

        Returns:
            List of key phrases
        """
        try:
            # First try the noun phrase extraction approach
            from nltk import pos_tag, word_tokenize
            from nltk.tokenize import sent_tokenize

            key_phrases = []

            # Process by sentence to maintain context
            for sentence in sent_tokenize(text):
                # Tag words with parts of speech
                tagged_tokens = pos_tag(word_tokenize(sentence))

                i = 0
                while i < len(tagged_tokens):
                    # Look for noun phrase patterns:
                    # 1. Optional determiners/adjectives + Nouns
                    # 2. Proper noun sequences

                    # Start with adjectives, determiners or nouns
                    if (tagged_tokens[i][1].startswith('JJ') or
                        tagged_tokens[i][1].startswith('DT') or
                            tagged_tokens[i][1].startswith('NN')):

                        start_idx = i
                        has_noun = tagged_tokens[i][1].startswith('NN')

                        # Continue while we see adjectives, determiners, or nouns
                        i += 1
                        while i < len(tagged_tokens) and (
                            tagged_tokens[i][1].startswith('JJ') or
                            tagged_tokens[i][1].startswith('DT') or
                            tagged_tokens[i][1].startswith('NN') or
                            # Include some prepositions
                            tagged_tokens[i][1] == 'IN'
                        ):
                            if tagged_tokens[i][1].startswith('NN'):
                                has_noun = True
                            i += 1

                        # Valid phrase must have at least one noun and be 2+ words
                        if has_noun and (i - start_idx) >= 2:
                            phrase = ' '.join(
                                token[0] for token in tagged_tokens[start_idx:i])
                            if phrase.lower() not in [kp.lower() for kp in key_phrases]:
                                key_phrases.append(phrase)
                    else:
                        i += 1

            # Extract specialized format phrases (like "X-based", "X-driven")
            specialized_patterns = [
                r'\b[A-Za-z]+-based\b',
                r'\b[A-Za-z]+-driven\b',
                r'\b[A-Za-z]+-oriented\b',
                r'\b[A-Za-z]+ [A-Za-z]+tion\b',  # action phrases
            ]

            for pattern in specialized_patterns:
                special_phrases = re.findall(pattern, text)
                for phrase in special_phrases:
                    if phrase not in key_phrases:
                        key_phrases.append(phrase)

            # Return most common phrases
            phrase_counter = Counter(key_phrases)
            return [phrase for phrase, _ in phrase_counter.most_common(max_phrases)]

        except Exception as e:
            print(f"Error extracting phrases: {str(e)}")
            # Fallback to simple ngram extraction
            words = [w.lower() for w in word_tokenize(text)
                     if w.lower() not in self.stop_words and len(w) >= 3]
            bigrams = [' '.join(words[i:i+2]) for i in range(len(words)-1)]
            return list(set(bigrams))[:max_phrases]

    def preprocess_text(self, text: str, use_stemming: bool = False) -> str:
        """
        Preprocess text for BM25 indexing

        Args:
            text: The input text to preprocess
            use_stemming: Whether to apply stemming (default: False)

        Returns:
            Preprocessed text
        """
        try:
            tokens = word_tokenize(text.lower())
            if use_stemming:
                processed = [
                    self.stemmer.stem(t) for t in tokens
                    if t.isalpha() and
                    t not in self.stop_words and
                    len(t) >= self.min_word_length
                ]
            else:
                processed = [
                    t for t in tokens
                    if t.isalpha() and
                    t not in self.stop_words and
                    len(t) >= self.min_word_length
                ]
            return ' '.join(processed)
        except Exception as e:
            print(f"Error preprocessing text: {str(e)}")
            return text.lower()

    def extract_all_concepts(self, text: str) -> dict:
        """
        Extract all types of concepts from text (comprehensive)

        Args:
            text: Input text

        Returns:
            Dictionary with different types of extracted concepts
        """
        results = {}

        # Extract single-word keywords
        results['keywords'] = self.extract_keywords(text)

        # Extract multi-word phrases
        results['key_phrases'] = self.extract_key_phrases(text)

        # Extract abbreviations and acronyms
        results['abbreviations'] = re.findall(r'\b[A-Z][A-Z0-9]+\b', text)

        # Extract proper nouns (names, organizations, etc.)
        try:
            from nltk import pos_tag, word_tokenize
            tagged_words = pos_tag(word_tokenize(text))
            proper_nouns = [word for word,
                            tag in tagged_words if tag.startswith('NNP')]
            results['proper_nouns'] = proper_nouns
        except:
            results['proper_nouns'] = []

        return results


class ContextualRetriever:
    """Class to handle contextual retrieval and enrichment of document chunks"""

    def __init__(self, llm, vectorstore, bm25_retriever):
        self.llm = llm
        self.vectorstore = vectorstore
        self.bm25_retriever = bm25_retriever
        self.keyword_processor = KeywordProcessor()

        # Define prompts for both enrichment and retrieval
        self.enrichment_prompt = PromptTemplate(
            template="""<document>{document}</document>
            <chunk>{chunk}</chunk>
            Provide 1-2 sentences explaining this chunk's context within the document:""",
            input_variables=["document", "chunk"]
        )

        self.retrieval_prompt = PromptTemplate(
            template="""Document: {document}
            Question: {question}
            Chunk: {chunk}
            Score this chunk's relevance (0-1) and explain why:""",
            input_variables=["document", "question", "chunk"]
        )

    def _analyze_query(self, question: str) -> dict:
        """Determine query characteristics to guide retrieval"""
        analysis = {
            'type': 'fact',  # fact, opinion, summary, etc.
            'entities': [],
            'keywords': []
        }

        # Simple heuristic - could be enhanced with LLM
        question_lower = question.lower()
        if 'what is' in question_lower or 'who is' in question_lower:
            analysis['type'] = 'definition'
        elif 'how to' in question_lower:
            analysis['type'] = 'process'
        elif 'why' in question_lower:
            analysis['type'] = 'explanation'

        # Extract entities and keywords (simple version)
        analysis['keywords'] = [word for word in question.split()
                                if word.lower() not in ['what', 'how', 'why', 'when']]

        return analysis

    def hybrid_search(self, question: str, k: int = 5) -> List[Document]:
        """Improved hybrid search with consistent scoring and deduplication"""
        try:
            # 1. Retrieve documents from both methods
            vector_docs = self.vectorstore.similarity_search(question, k=k*3)
            bm25_docs = self.bm25_retriever.retrieve(question, k=k*3)

            # 2. Create unified scoring dictionary with content as key
            doc_dict = {}

            # Process vector docs first
            for doc in vector_docs:
                if not isinstance(doc, Document):
                    continue

                content = doc.page_content
                if content not in doc_dict:
                    doc_dict[content] = {
                        'doc': doc,
                        'vector_score': doc.metadata.get('similarity_score', 0),
                        'bm25_score': 0  # Initialize BM25 score
                    }
                else:
                    # Keep the highest vector score if we see duplicates
                    doc_dict[content]['vector_score'] = max(
                        doc_dict[content]['vector_score'],
                        doc.metadata.get('similarity_score', 0)
                    )

            # Process BM25 docs
            for doc in bm25_docs:
                if not isinstance(doc, Document):
                    continue

                content = doc.page_content
                if content not in doc_dict:
                    doc_dict[content] = {
                        'doc': doc,
                        'vector_score': 0,  # Initialize vector score
                        'bm25_score': doc.metadata.get('bm25_score', 0)
                    }
                else:
                    # Keep the highest BM25 score if we see duplicates
                    doc_dict[content]['bm25_score'] = max(
                        doc_dict[content]['bm25_score'],
                        doc.metadata.get('bm25_score', 0)
                    )

            # 3. Normalize and combine scores
            if not doc_dict:
                return []

            # Get max scores for normalization
            max_vector = max(info['vector_score']
                             for info in doc_dict.values()) or 1
            max_bm25 = max(info['bm25_score']
                           for info in doc_dict.values()) or 1

            # Calculate combined scores
            scored_docs = []
            for content, info in doc_dict.items():
                # Normalize scores
                norm_vector = info['vector_score'] / max_vector
                norm_bm25 = info['bm25_score'] / max_bm25

                # Weighted combination (adjust weights as needed)
                combined_score = (norm_vector * 0.4) + (norm_bm25 * 0.6)

                # Create enriched document
                metadata = info['doc'].metadata.copy()
                metadata.update({
                    'vector_score': norm_vector,
                    'bm25_score': norm_bm25,
                    'combined_score': combined_score,
                    'relevance_percentage': int(combined_score * 100)
                })

                scored_docs.append((
                    Document(page_content=content, metadata=metadata),
                    combined_score
                ))

            # 4. Sort by combined score (descending)
            scored_docs.sort(key=lambda x: x[1], reverse=True)

            # 5. Apply contextual analysis to top k unique documents
            top_docs = [doc for doc, score in scored_docs[:k]]
            return self._apply_contextual_analysis(question, top_docs)

        except Exception as e:
            logger.error(f"Error in hybrid search: {str(e)}")
            return []

    def _apply_contextual_analysis(self, question: str, docs: List[Document]) -> List[Document]:
        """Apply contextual analysis to retrieved documents"""
        processed_docs = []

        for doc in docs:
            try:
                # Get full document context from metadata
                full_context = doc.metadata.get(
                    "full_document", doc.page_content)

                # Run contextual analysis
                analysis_prompt = """Given this document context and user question, please:
                                1. Provide a relevance score between 0 and 1 (where 1 is most relevant)
                                2. Explain in one sentence why this is relevant
                                3. List key connections

                                DOCUMENT CONTEXT:
                                {document_context}

                                USER QUESTION:
                                {question}

                                CHUNK CONTENT:
                                {chunk_content}

                                Respond in this exact format:
                                Score: [0-1]
                                Explanation: [your explanation]
                                Connections: [comma-separated list]"""

                analysis = self.llm.invoke(analysis_prompt.format(
                    document_context=full_context,
                    question=question,
                    chunk_content=doc.page_content
                )).content

                # Parse the analysis response
                score = 0.5  # Default score if parsing fails
                explanation = ""
                connections = []

                try:
                    # Extract score
                    score_line = next(line for line in analysis.split(
                        '\n') if line.startswith('Score:'))
                    score = float(score_line.split(':')[1].strip())
                    score = max(0, min(1, score))  # Clamp between 0-1

                    # Extract explanation
                    explanation_line = next(line for line in analysis.split(
                        '\n') if line.startswith('Explanation:'))
                    explanation = explanation_line.split(':', 1)[1].strip()

                    # Extract connections
                    connections_line = next(line for line in analysis.split(
                        '\n') if line.startswith('Connections:'))
                    connections = [c.strip() for c in connections_line.split(':', 1)[
                        1].split(',')]
                except (StopIteration, ValueError, IndexError) as e:
                    logger.warning(
                        f"Couldn't fully parse analysis: {str(e)}. Using defaults.")

                # Update document metadata with contextual analysis
                enriched_metadata = doc.metadata.copy()
                enriched_metadata.update({
                    "contextual_score": float(score),
                    "contextual_explanation": explanation,
                    "contextual_connections": connections,
                    "retrieval_method": "hybrid"
                })

                # Create enriched document
                enriched_doc = Document(
                    page_content=doc.page_content,
                    metadata=enriched_metadata
                )
                processed_docs.append(enriched_doc)

            except Exception as e:
                logger.error(f"Error processing document: {str(e)}")
                continue

        return processed_docs

    def contextual_search(self, question: str, k: int = 5) -> List[Document]:
        """Use hybrid search by default"""
        return self.hybrid_search(question, k)

    def enrich_chunk(self, doc: Document) -> Document:
        """Enhanced contextual enrichment with better error handling"""
        try:
            if not isinstance(doc, Document):
                raise ValueError("Input must be a Document object")

            # Get full document context
            full_doc = doc.metadata.get("full_document", "")
            if not full_doc:
                logger.warning(
                    "No full document context available for enrichment")
                return doc

            # Skip enrichment if chunk is too small (may not have enough context)
            if len(doc.page_content) < 50:
                logger.debug("Skipping enrichment for very small chunk")
                return doc

            try:
                # Run enrichment chain
                chain = self.enrichment_prompt | self.llm | StrOutputParser()
                context = chain.invoke({
                    "document": full_doc,
                    "chunk": doc.page_content
                }).strip()

                if not context:
                    raise ValueError("Empty context returned from enrichment")

                # Create new metadata
                new_metadata = doc.metadata.copy()
                new_metadata.update({
                    "contextual_summary": context,
                    "enrichment_time": time.time()
                })

                return Document(
                    page_content=f"CONTEXT: {context}\nCONTENT: {doc.page_content}",
                    metadata=new_metadata
                )

            except Exception as e:
                logger.error(f"Error in enrichment chain: {str(e)}")
                return doc  # Return original doc if enrichment fails

        except Exception as e:
            logger.error(f"Error in enrich_chunk: {str(e)}")
            return doc  # Fallback to original document


class ContextualRAGApplication:
    """Enhanced RAG application with contextual embedding"""

    def __init__(self, persist_directory: str = "./chroma_db_docx"):
        """Initialize the Contextual RAG application"""
        self.persist_directory = persist_directory
        self.embeddings = None
        self.llm = None
        self.vectorstore = None
        self.retriever = None
        self.qa_chain = None
        self.contextual_enricher = None
        self.memory = ConversationBufferWindowMemory(k=5, return_messages=True)
        self.bm25_retriever = None

        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,  # Reduced from 1000
            chunk_overlap=150,  # Reduced from 200
            length_function=len,
            add_start_index=True,
            separators=["\n\n", "\n", " ", ""]  # Explicit separators
        )

        # Load environment variables
        load_dotenv()

        try:
            # Initialize components in proper order
            self._initialize_azure_components()
            self._initialize_vectorstore()

            # Now that vectorstore is initialized, create BM25 retriever
            self.bm25_retriever = BM25Retriever(self.vectorstore)
            if self.bm25_retriever.bm25 is None:
                logger.warning(
                    "BM25 retriever disabled - no documents available")

            # Initialize contextual components
            self.contextual_enricher = ContextualRetriever(
                self.llm, self.vectorstore, self.bm25_retriever)
            self._initialize_rag_pipeline()

            logger.info("Contextual RAG application initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize BM25 retriever: {str(e)}")
            self.bm25_retriever = None
            raise

    def _initialize_azure_components(self):
        """Initialize Azure OpenAI components"""
        try:
            # Get environment variables with fallbacks
            embedding_deployment = os.getenv(
                "AZURE_OPENAI_EMBEDDING_DEPLOYMENT")
            chat_deployment = os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT")
            api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2023-05-15")
            endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
            api_key = os.getenv("AZURE_OPENAI_API_KEY")

            # Check for required configuration
            if not all([embedding_deployment, chat_deployment, endpoint, api_key]):
                missing = []
                if not embedding_deployment:
                    missing.append("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")
                if not chat_deployment:
                    missing.append("AZURE_OPENAI_CHAT_DEPLOYMENT")
                if not endpoint:
                    missing.append("AZURE_OPENAI_ENDPOINT")
                if not api_key:
                    missing.append("AZURE_OPENAI_API_KEY")

                raise ValueError(
                    f"Missing required Azure OpenAI configuration: {', '.join(missing)}")

            logger.info(
                f"Initializing Azure OpenAI with: Embedding deployment: {embedding_deployment}, Chat deployment: {chat_deployment}")

            # Initialize embeddings
            try:
                self.embeddings = AzureOpenAIEmbeddings(
                    azure_deployment=embedding_deployment,
                    openai_api_version=api_version,
                    azure_endpoint=endpoint,
                    api_key=api_key
                )
                # Test the embeddings with a simple query to catch errors early
                test_result = self.embeddings.embed_query("Test embedding")
                if not test_result or len(test_result) < 1:
                    raise ValueError(
                        "Embedding test failed - returned empty result")
                logger.info(
                    f"Embeddings successfully initialized and tested ({len(test_result)} dimensions)")
            except Exception as e:
                error_msg = str(e)
                if "OperationNotSupported" in error_msg:
                    raise ValueError(
                        f"The model in deployment '{embedding_deployment}' doesn't support embeddings. "
                        f"Please use a text-embedding model like 'text-embedding-ada-002'. Error: {error_msg}"
                    )
                else:
                    raise ValueError(
                        f"Failed to initialize embeddings: {error_msg}"
                    )

            # Initialize LLM
            try:
                self.llm = AzureChatOpenAI(
                    azure_deployment=chat_deployment,
                    openai_api_version=api_version,
                    azure_endpoint=endpoint,
                    api_key=api_key,
                    temperature=0.1,  # Low temperature for factual responses
                    max_tokens=1000
                )
                logger.info(f"Chat model successfully initialized")
            except Exception as e:
                raise ValueError(f"Failed to initialize chat model: {str(e)}")

            logger.info("Azure OpenAI components initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize Azure components: {str(e)}")
            raise

    def _initialize_contextual_enricher(self):
        """Initialize the contextual enrichment component"""
        try:
            self.contextual_enricher = ContextualRetriever(
                self.llm, self.vectorstore)
            logger.info("Contextual enricher initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize contextual enricher: {str(e)}")
            raise

    def _initialize_vectorstore(self):
        """Initialize or load the vector store"""
        try:
            # Check if persist directory exists
            if os.path.exists(self.persist_directory) and os.listdir(self.persist_directory):
                # Load existing vector store
                logger.info(
                    f"Loading vector store from {self.persist_directory}")
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                logger.info(
                    f"Loaded existing vector store with {self.vectorstore._collection.count()} documents")
            else:
                # Create directory if it doesn't exist
                os.makedirs(self.persist_directory, exist_ok=True)

                # Create new vector store
                logger.info("Creating new vector store")
                self.vectorstore = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
                self.vectorstore.persist()
                logger.info("Created new empty vector store")

            # Verify the vectorstore is properly initialized
            if not hasattr(self.vectorstore, '_collection'):
                raise RuntimeError(
                    "Vector store failed to initialize properly - missing _collection attribute")

        except Exception as e:
            logger.error(f"Failed to initialize vector store: {str(e)}")
            self.vectorstore = None  # Ensure it's None if initialization fails
            raise

    def _initialize_rag_pipeline(self):
        """Initialize RAG pipeline with hybrid retrieval"""
        try:
            # Update the prompt to include retrieval method info
            prompt_template = """You are a document intelligence assistant. Use the following context to answer:

            CONTEXTUAL ANALYSIS:
            {contextual_analysis}

            CHUNK CONTENT:
            {context}

            CHAT HISTORY:
            {chat_history}

            QUESTION: {question}

            Guidelines:
            1. First analyze the contextual relevance information
            2. Focus on answers that synthesize information across chunks
            3. For factual questions, verify against multiple chunks when possible
            4. If unsure, say "I couldn't find definitive information about this in the documents"
            5. Consider that some results came from semantic search while others came from keyword matching
            """

            PROMPT = PromptTemplate(
                template=prompt_template,
                input_variables=["contextual_analysis",
                                 "context", "question", "chat_history"]
            )

            # Create a runnable for hybrid retrieval
            def hybrid_retriever(input_dict: Dict[str, Any]) -> Dict[str, Any]:
                """Retrieve documents with hybrid search and contextual analysis"""
                try:
                    question = input_dict["question"]
                    docs = self.contextual_enricher.hybrid_search(question)

                    contextual_analysis = "\n\n".join(
                        f"Document {i+1} (Relevance: {doc.metadata.get('contextual_score', 0):.0%}, Method: {doc.metadata.get('retrieval_method', 'unknown')}):\n"
                        f"Explanation: {doc.metadata.get('contextual_explanation', '')}\n"
                        f"Connections: {', '.join(doc.metadata.get('contextual_connections', []))}"
                        for i, doc in enumerate(docs)
                    )

                    context_content = "\n\n".join(
                        f"Document {i+1}:\n{doc.page_content}"
                        for i, doc in enumerate(docs))

                    return {
                        "contextual_analysis": contextual_analysis,
                        "context": context_content,
                        "question": question,
                        "chat_history": input_dict.get("chat_history", "")
                    }
                except Exception as e:
                    logger.error(f"Error in hybrid retriever: {str(e)}")
                    return {
                        "contextual_analysis": "Error retrieving contextual information",
                        "context": "",
                        "question": question,
                        "chat_history": input_dict.get("chat_history", "")
                    }

            # Build the full chain with hybrid retrieval
            self.qa_chain = (
                RunnablePassthrough()
                | hybrid_retriever
                | PROMPT
                | self.llm
                | StrOutputParser()
            )

            logger.info("Hybrid RAG pipeline initialized successfully")

        except Exception as e:
            logger.error(f"Error initializing RAG pipeline: {str(e)}")
            raise

    def _get_chat_history(self) -> str:
        """Get formatted chat history from memory"""
        messages = self.memory.load_memory_variables({}).get("history", [])
        formatted_messages = []

        for message in messages:
            if hasattr(message, "content"):
                role = "Human" if message.type == "human" else "Assistant"
                formatted_messages.append(f"{role}: {message.content}")

        return "\n".join(formatted_messages)

    def add_uploaded_files(self, files) -> str:
        """Process and add uploaded files to the vector store"""
        try:
            if not files:
                return "No files provided"

            documents = []
            processed_files = 0
            failed_files = 0

            for file in files:
                file_name = "unknown"
                try:
                    # Handle different file input types
                    if isinstance(file, tuple) and len(file) == 2:
                        file_name, file_content = file
                        file_bytes = file_content
                    elif hasattr(file, 'name') and hasattr(file, 'read'):
                        file_name = file.name
                        file.seek(0)
                        file_bytes = file.read()
                    else:
                        file_name = f"upload_{time.time()}.docx"
                        file_bytes = file

                    # Validate file type
                    if not file_name.lower().endswith('.docx'):
                        logger.warning(f"Skipping non-DOCX file: {file_name}")
                        continue

                    # Create temp directory if it doesn't exist
                    temp_dir = "./temp_uploads"
                    os.makedirs(temp_dir, exist_ok=True)

                    # Create temp file path
                    file_hash = hashlib.md5(file_name.encode()).hexdigest()[:8]
                    clean_name = re.sub(
                        r'[^\w.-]', '_', os.path.basename(file_name))
                    temp_path = os.path.join(
                        temp_dir, f"upload_{file_hash}_{clean_name}")

                    # Write to temp file
                    with open(temp_path, "wb") as f:
                        f.write(file_bytes)

                    # Verify file is valid DOCX
                    try:
                        doc = DocxDocument(temp_path)
                        if not doc.paragraphs:
                            raise ValueError("No paragraphs found in document")
                    except Exception as e:
                        logger.warning(
                            f"Invalid DOCX file {file_name}: {str(e)}")
                        failed_files += 1
                        continue

                    # Load document
                    # Load and process document
                    # Load and process document
                    loader = DocxLoader(temp_path)
                    loaded_docs = loader.load()

                    for doc in loaded_docs:
                        if doc.page_content.strip():
                            # Apply metadata filtering immediately using our standard function
                            filtered_doc = safe_filter_metadata(doc)
                            if isinstance(filtered_doc, Document) and filtered_doc.page_content.strip():
                                documents.append(filtered_doc)

                except Exception as e:
                    logger.error(f"Error processing file: {str(e)}")
                    continue

            if not documents:
                return "No valid documents were extracted"

            # Process documents through the standardized pipeline
            processed_chunks = self._process_documents(documents)

            if not processed_chunks:
                return "No valid chunks created"

            # Store with additional metadata validation
            try:
                # Double-check metadata before storing using our standard function
                safe_chunks = []
                for chunk in processed_chunks:
                    filtered = safe_filter_metadata(chunk)
                    if isinstance(filtered, Document) and filtered.page_content.strip():
                        safe_chunks.append(filtered)

                if safe_chunks:
                    self.vectorstore.add_documents(safe_chunks)
                    self.vectorstore.persist()

                    # Reinitialize retrievers
                    if self.bm25_retriever:
                        self.bm25_retriever._initialize_bm25()

                    return f"Successfully stored {len(safe_chunks)} document chunks"
                else:
                    return "No valid chunks after final filtering"

            except Exception as e:
                logger.error(f"Error storing documents: {str(e)}")
                return f"Error storing documents: {str(e)}"

        except Exception as e:
            logger.error(f"Error in add_uploaded_files: {str(e)}")
            return f"Error processing files: {str(e)}"

    def _process_documents(self, documents: List[Document]) -> List[Document]:
        """Process documents with proper metadata filtering and contextual enrichment"""
        processed_chunks = []
        keyword_processor = KeywordProcessor()

        for doc in documents:
            if not isinstance(doc, Document) or not doc.page_content.strip():
                continue

            try:
                # First apply metadata filtering to the original document
                filtered_doc = safe_filter_metadata(doc)
                metadata = filtered_doc.metadata.copy()

                # Ensure required metadata fields
                metadata.update({
                    "source": metadata.get("source", "unknown"),
                    "filename": metadata.get("filename", "unknown"),
                    "full_document": metadata.get("full_document", filtered_doc.page_content),
                    "processing_time": time.time()
                })

                # Extract keywords if not present (with filtering)
                if "keywords" not in metadata:
                    keywords = keyword_processor.extract_keywords(
                        filtered_doc.page_content)
                    metadata["keywords"] = ", ".join(
                        keywords) if isinstance(keywords, list) else ""

                # Add preprocessed text if not present
                if "preprocessed_text" not in metadata:
                    preprocessed = keyword_processor.preprocess_text(
                        filtered_doc.page_content)
                    metadata["preprocessed_text"] = preprocessed if isinstance(
                        preprocessed, str) else ""

                # Split document into chunks
                chunks = self.text_splitter.split_documents([filtered_doc])

                for chunk in chunks:
                    if not isinstance(chunk, Document) or not chunk.page_content.strip():
                        continue

                    try:
                        # Prepare chunk metadata with additional filtering
                        chunk_metadata = metadata.copy()
                        chunk_metadata.update({
                            "original_content": chunk.page_content,
                            "is_full_document": False,
                            "chunk_length": len(chunk.page_content)
                        })

                        # Apply final metadata filtering
                        filtered_chunk_metadata = safe_filter_metadata(
                            chunk_metadata)
                        if isinstance(filtered_chunk_metadata, dict):
                            final_metadata = filtered_chunk_metadata
                        else:
                            final_metadata = filtered_chunk_metadata.metadata

                        # Create base chunk with filtered metadata
                        base_chunk = Document(
                            page_content=chunk.page_content,
                            metadata=final_metadata
                        )

                        # Apply contextual enrichment if available
                        if hasattr(self, 'contextual_enricher') and self.contextual_enricher:
                            try:
                                enriched_chunk = self.contextual_enricher.enrich_chunk(
                                    base_chunk)
                                if enriched_chunk and enriched_chunk.page_content.strip():
                                    # Final metadata filtering before adding
                                    final_chunk = safe_filter_metadata(
                                        enriched_chunk)
                                    if isinstance(final_chunk, Document) and final_chunk.page_content.strip():
                                        processed_chunks.append(final_chunk)
                                        continue
                            except Exception as enrich_error:
                                logger.error(
                                    f"Contextual enrichment failed: {str(enrich_error)}")

                        # Fallback: add base chunk if it's a Document with content
                        if isinstance(base_chunk, Document) and base_chunk.page_content.strip():
                            processed_chunks.append(base_chunk)

                    except Exception as chunk_error:
                        logger.error(
                            f"Error processing chunk: {str(chunk_error)}")
                        continue

            except Exception as doc_error:
                logger.error(f"Error processing document: {str(doc_error)}")
                continue

        return processed_chunks

    def query(self, question: str, add_to_memory: bool = True) -> Dict[str, Any]:
        """Query the contextual RAG pipeline"""
        try:
            # Check if components are properly initialized
            if self.vectorstore is None:
                return {
                    "result": "Document system is not properly initialized. Please try resetting the document store.",
                    "source_documents": [],
                    "query_time": None
                }

            if not hasattr(self.vectorstore, '_collection'):
                return {
                    "result": "Vector store is not properly configured. Please try resetting the document store.",
                    "source_documents": [],
                    "query_time": None
                }

            # Check if the vector store is empty
            try:
                if self.vectorstore._collection.count() == 0:
                    return {
                        "result": "No documents have been added to the system. Please upload some documents first.",
                        "source_documents": [],
                        "query_time": None
                    }
            except Exception as e:
                logger.error(f"Error checking vector store count: {str(e)}")
                return {
                    "result": "Error checking document store status.",
                    "source_documents": [],
                    "query_time": None
                }

            # Get start time
            start_time = time.time()

            # Get the answer from the QA chain
            try:
                answer = self.qa_chain.invoke({
                    "question": question,
                    "chat_history": self._get_chat_history()
                })
            except Exception as e:
                logger.error(f"Error in QA chain: {str(e)}")
                answer = f"Error generating answer: {str(e)}"

            # Get the source documents with contextual information
            try:
                source_docs = []
                if self.contextual_enricher is not None:
                    source_docs = self.contextual_enricher.contextual_search(
                        question)
            except Exception as e:
                logger.error(f"Error retrieving source documents: {str(e)}")
                source_docs = []

            # Calculate query time
            query_time = time.time() - start_time

            # Add to memory if requested
            if add_to_memory and self.memory is not None:
                try:
                    self.memory.save_context(
                        {"input": question},
                        {"output": answer}
                    )
                except Exception as e:
                    logger.error(f"Error saving to memory: {str(e)}")

            return {
                "result": answer,
                "source_documents": source_docs,
                "query_time": f"{query_time:.2f} seconds"
            }
        except Exception as e:
            logger.error(f"Unexpected error in query: {str(e)}", exc_info=True)
            return {
                "result": f"Unexpected error processing your request: {str(e)}",
                "source_documents": [],
                "query_time": None
            }

    def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about the documents in the vector store"""
        try:
            # Get document count
            doc_count = self.vectorstore._collection.count()

            # Get unique sources
            if doc_count > 0:
                docs = self.vectorstore._collection.get()
                metadatas = docs["metadatas"]
                unique_sources = set()

                # Collect topic statistics
                all_topics = []
                all_entities = []

                for metadata in metadatas:
                    if "source" in metadata:
                        unique_sources.add(metadata["source"])

                    # Collect topics and entities for statistics
                    if "topics" in metadata and isinstance(metadata["topics"], list):
                        all_topics.extend(metadata["topics"])
                    if "entities" in metadata and isinstance(metadata["entities"], list):
                        all_entities.extend(metadata["entities"])

                # Get top topics and entities
                top_topics = self._get_top_items(all_topics, 10)
                top_entities = self._get_top_items(all_entities, 10)

            else:
                unique_sources = set()
                top_topics = []
                top_entities = []

            return {
                "document_chunks": doc_count,
                "unique_files": len(unique_sources),
                "files": list(unique_sources),
                "top_topics": top_topics,
                "top_entities": top_entities
            }
        except Exception as e:
            logger.error(f"Error getting document stats: {str(e)}")
            return {
                "document_chunks": 0,
                "unique_files": 0,
                "files": [],
                "top_topics": [],
                "top_entities": []
            }

    def _get_top_items(self, items: List[str], limit: int = 10) -> List[Tuple[str, int]]:
        """Get top items by frequency"""
        from collections import Counter
        if not items:
            return []

        counter = Counter(items)
        return counter.most_common(limit)

    def clear_memory(self) -> None:
        """Clear the conversation memory"""
        self.memory.clear()
        logger.info("Conversation memory cleared")

    def reset_vectorstore(self) -> str:
        """Reset the vector store, removing all documents"""
        try:
            # Close existing connections first
            if self.vectorstore:
                try:
                    self.vectorstore._client = None
                    self.vectorstore._collection = None
                except Exception as e:
                    logger.warning(f"Error cleaning up vectorstore: {str(e)}")

            # Delete the persist directory
            import shutil
            if os.path.exists(self.persist_directory):
                for _ in range(3):  # Retry up to 3 times
                    try:
                        shutil.rmtree(self.persist_directory)
                        break
                    except PermissionError as e:
                        logger.warning(f"Retrying delete: {str(e)}")
                        time.sleep(0.5)  # Wait briefly before retrying
                    except Exception as e:
                        logger.error(f"Error deleting directory: {str(e)}")
                        raise

            # Reinitialize the vector store
            self._initialize_vectorstore()

            logger.info("Vector store reset successfully")
            return "✅ Vector store reset successfully. All documents have been removed."
        except Exception as e:
            logger.error(f"Error resetting vector store: {str(e)}")
            return f"❌ Error resetting vector store: {str(e)}"

    # def search_by_metadata(self, field: str, query: str) -> List[Document]:
        # """Search for documents with specific metadata field values"""
        # try:
        #     # Get all documents
        #     all_docs = self.vectorstore._collection.get()
        #     metadatas = all_docs["metadatas"]
        #     documents = all_docs["documents"]

        #     # Filter documents by metadata field
        #     results = []
        #     for i, metadata in enumerate(metadatas):
        #         # Check if the field exists and contains the query
        #         if field in metadata:
        #             field_value = metadata[field]

        #             # Handle different types of metadata fields
        #             if isinstance(field_value, list):
        #                 # For lists (like topics, entities)
        #                 if any(query.lower() in item.lower() for item in field_value):
        #                     results.append(Document(
        #                         page_content=documents[i],
        #                         metadata=metadata
        #                     ))
        #             elif isinstance(field_value, str):
        #                 # For string fields
        #                 if query.lower() in field_value.lower():
        #                     results.append(Document(
        #                         page_content=documents[i],
        #                         metadata=metadata
        #                     ))

        #     return results
        # except Exception as e:
        #     logger.error(f"Error searching by metadata: {str(e)}")
        #     return []


# Initialize the RAG application
rag_app = ContextualRAGApplication()

# Gradio Interface


def query_rag(question, history, show_sources, include_time, show_context):
    """Handle the question and return RAG response with contextual sources"""
    try:
        # Query the RAG pipeline
        result = rag_app.query(question)
        response = result["result"]

        # Add sources if requested
        if show_sources and "source_documents" in result:
            sources = "\n\n📄 **Document Sources (with Contextual Relevance):**\n"

            for i, doc in enumerate(result["source_documents"], 1):
                source_name = os.path.basename(
                    doc.metadata.get('source', 'Unknown'))
                score = doc.metadata.get("contextual_score", 0)
                score_percent = int(score * 100)

                sources += f"\n**Source {i}: {source_name} (Relevance: {score_percent}%)**\n"

                if show_context:
                    explanation = doc.metadata.get(
                        "contextual_explanation", "")
                    connections = doc.metadata.get(
                        "contextual_connections", [])

                    sources += f"*Why relevant:* {explanation}\n"
                    if connections:
                        sources += f"*Key connections:* {', '.join(connections)}\n"

                # Show the original content
                chunk_content = doc.metadata.get(
                    "original_content", doc.page_content)
                sources += f'"{chunk_content}"\n'

            response += sources

        # Add query time if requested
        if include_time and "query_time" in result:
            response += f"\n\n⏱️ Query time: {result['query_time']}"

        return response

    except Exception as e:
        logger.error(f"Error in query_rag: {str(e)}")
        return f"Error processing your request: {str(e)}"


def upload_files(files):
    """Handle DOCX file uploads"""
    if not files:
        return "No files uploaded"

    logger.info(f"Received {len(files)} files for upload")

    # Process files
    result = rag_app.add_uploaded_files(files)

    # Get updated stats
    stats = rag_app.get_document_stats()

    # Return status with stats
    return f"{result}\n\nCurrent document store contains {stats['document_chunks']} chunks from {stats['unique_files']} files."


def clear_conversation():
    """Clear the conversation memory"""
    rag_app.clear_memory()
    return "Conversation memory cleared"


def reset_documents():
    """Reset the document store"""
    result = rag_app.reset_vectorstore()

    # Clear memory as well
    rag_app.clear_memory()

    return result


def display_document_stats():
    """Display statistics about the documents"""
    stats = rag_app.get_document_stats()

    if stats["document_chunks"] == 0:
        return "No documents have been added to the system."

    response = f"📊 **Document Statistics**\n\n"
    response += f"- Total document chunks: {stats['document_chunks']}\n"
    response += f"- Unique files: {stats['unique_files']}\n\n"

    if stats["top_topics"]:
        response += "**Top Topics:**\n"
        for topic, count in stats["top_topics"]:
            response += f"- {topic} ({count})\n"
        response += "\n"

    if stats["top_entities"]:
        response += "**Top Entities:**\n"
        for entity, count in stats["top_entities"]:
            response += f"- {entity} ({count})\n"
        response += "\n"

    if stats["files"]:
        response += "**Files:**\n"
        for i, file in enumerate(stats["files"], 1):
            response += f"{i}. {os.path.basename(file)}\n"

    return response


# Create Gradio Interface
with gr.Blocks(theme=gr.themes.Soft()) as demo:
    gr.Markdown("""
    # 📄 Contextual Document Intelligence
    **Azure OpenAI-powered Contextual RAG application**
    
    This application uses contextual retrieval to enrich document chunks with additional information before indexing them.
    """)

    with gr.Tab("💬 Chat with Documents"):
        with gr.Row():
            with gr.Column(scale=4):
                chatbot = gr.ChatInterface(
                    fn=query_rag,
                    additional_inputs=[
                        gr.Checkbox(
                            label="Show document sources with content", value=True),
                        gr.Checkbox(label="Show query time", value=False),
                        gr.Checkbox(
                            label="Show contextual summaries", value=True)
                    ]
                )

            with gr.Column(scale=1):
                gr.Markdown("### Options")
                clear_memory_btn = gr.Button("Clear Conversation Memory")
                clear_memory_btn.click(
                    clear_conversation,
                    inputs=[],
                    outputs=[gr.Textbox(label="Status")]
                )

                stats_btn = gr.Button("Show Document Statistics")
                stats_output = gr.Textbox(label="Document Stats", lines=10)
                stats_btn.click(
                    display_document_stats,
                    inputs=[],
                    outputs=stats_output
                )

    with gr.Tab("📤 Upload Documents"):
        with gr.Row():
            with gr.Column(scale=3):
                file_input = gr.File(
                    label="Upload DOCX Files",
                    file_types=[".docx"],
                    file_count="multiple",
                    type="binary"
                )
                upload_output = gr.Textbox(label="Upload Status", lines=5)
                file_input.upload(
                    upload_files,
                    inputs=file_input,
                    outputs=upload_output
                )

            with gr.Column(scale=2):
                gr.Markdown("""
                **Instructions:**
                1. Upload one or more DOCX files
                2. Switch to Chat tab to ask questions
                3. Toggle "Show document sources" to see references with content and relevance scores
                """)

                reset_btn = gr.Button("Reset Document Store", variant="stop")
                reset_output = gr.Textbox(label="Reset Status")
                reset_btn.click(
                    reset_documents,
                    inputs=[],
                    outputs=reset_output
                )

if __name__ == "__main__":
    try:
        # Launch the Gradio interface
        demo.launch(
            server_name="0.0.0.0",
            server_port=7860,
            show_error=True
        )
    except Exception as e:
        print(f"Error launching application: {str(e)}")
        print("Check the logs in 'rag_app.log' for more details.")
